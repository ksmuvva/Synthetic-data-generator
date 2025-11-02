"""
Fully NLP-based CLI - All interactions through natural language chat.
"""

import asyncio
import csv
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List

import typer
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Prompt
from rich.markdown import Markdown
from rich.table import Table
from dotenv import load_dotenv

from synth_agent import __version__
from synth_agent.core.config import Config

# Create single chat-based app
app = typer.Typer(
    name="synth-agent",
    help="🤖 AI-Powered Data Generator - Fully NLP Chat Interface",
    add_completion=False,
)

console = Console()


class ConversationContext:
    """Maintains conversation history and context."""

    def __init__(self):
        self.history: List[Dict[str, str]] = []
        self.last_generated_file: Optional[str] = None
        self.user_preferences: Dict[str, Any] = {
            "default_format": "csv",
            "default_rows": 10,
            "template_file": None  # Path to template file for pattern-based generation
        }
    
    def add_message(self, role: str, content: str):
        """Add message to history."""
        self.history.append({
            "role": role,
            "content": content,
            "timestamp": datetime.now().isoformat()
        })
    
    def get_context_summary(self) -> str:
        """Get summary of recent context."""
        if not self.history:
            return "No previous conversation."
        
        recent = self.history[-5:]  # Last 5 exchanges
        summary = []
        for msg in recent:
            role = msg["role"].capitalize()
            content = msg["content"][:100] + "..." if len(msg["content"]) > 100 else msg["content"]
            summary.append(f"{role}: {content}")
        
        return "\n".join(summary)


def print_welcome() -> None:
    """Print enhanced welcome banner."""
    welcome_text = f"""
# 🤖 Synthetic Data Generator v{__version__}

**Fully Natural Language Interface** - Just chat with the AI!

No commands to remember - just tell me what you need in plain English.

## What can I help you with?

💾 **Generate Data**: "create 100 customer records", "make JSON for API testing"
📊 **Analyze Data**: "analyze customers.csv", "show me patterns in sales data"
⚙️  **Configure**: "change output format to JSON", "set default rows to 50"
📋 **Template Files**: "use customers.csv as template", "set template to data.json"
📁 **File Operations**: "list my data files", "delete old_data.csv"
💡 **Help & Examples**: "show me examples", "how do I generate time series?"
🔍 **Ask Anything**: Questions about formats, data types, best practices

Just type naturally and I'll understand!
"""
    console.print(Panel(Markdown(welcome_text), border_style="cyan", title="Welcome"))


def print_examples():
    """Show example interactions."""
    examples = Table(title="💡 Example Conversations", show_header=True, header_style="bold magenta")
    examples.add_column("You Say", style="green", width=40)
    examples.add_column("I Understand", style="cyan", width=40)
    
    examples.add_row(
        "create 50 customer records with emails",
        "✓ Generate CSV with 50 customers including email field"
    )
    examples.add_row(
        "I need JSON data for testing my API",
        "✓ Generate JSON format data for API testing"
    )
    examples.add_row(
        "make an Excel file with sales data",
        "✓ Generate Excel (.xlsx) with sales information"
    )
    examples.add_row(
        "create a PDF document with customer bios",
        "✓ Generate PDF in paragraph format (natural document)"
    )
    examples.add_row(
        "create a PDF table with employee data",
        "✓ Generate PDF with table/spreadsheet layout"
    )
    examples.add_row(
        "can you analyze the patterns in my data?",
        "✓ Analyze file for patterns and insights"
    )
    examples.add_row(
        "change my default format to JSON",
        "✓ Update preferences to use JSON by default"
    )
    examples.add_row(
        "show me what files I have",
        "✓ List all generated data files"
    )
    
    console.print(examples)
    console.print()


async def classify_intent(client, user_input: str, context: ConversationContext) -> Dict[str, Any]:
    """Use Claude to classify user intent and extract parameters with intelligent reasoning."""

    system_prompt = """You are an AI assistant for a synthetic data generation tool with advanced reasoning capabilities.
Classify user requests into these intents:

1. GENERATE - User wants to create synthetic data
2. ANALYZE - User wants to analyze existing data
3. CONFIGURE - User wants to change settings/preferences
4. FILE_OPS - User wants to list/delete/manage files
5. HELP - User needs help or examples
6. QUESTION - User has a general question
7. CHAT - General conversation/greeting

Extract parameters based on intent:
- For GENERATE: data_description, rows, format (csv/json/excel/parquet/xml/txt/pdf/docx/word), fields, doc_style (paragraph/table - default paragraph)
- For ANALYZE: filename, analysis_type
- For CONFIGURE: setting_name, value
  - Template-related: "use X as template", "set template to Y" → setting_name="template_file", value="<file_path>"
  - Format: "use JSON format", "default to CSV" → setting_name="format", value="json/csv/etc"
  - Rows: "default 100 rows", "set rows to 50" → setting_name="rows", value=<number>
- For FILE_OPS: operation (list/delete/rename), filename

CRITICAL FORMAT DETECTION RULES - Apply intelligent reasoning:
1. **Explicit mentions**: "JSON", "CSV", "Excel", "PDF", "Word", "XML", "Parquet" → use that format
2. **Context clues**:
   - "API", "REST", "endpoint", "nested" → format="json"
   - "spreadsheet", "rows and columns", "table" → format="excel" or "csv"
   - "document", "paragraph", "essay", "article", "report" → format="pdf" or "docx"
   - "big data", "analytics", "data lake" → format="parquet"
   - "list", "simple text" → format="txt"
3. **Document style detection**:
   - For PDF/DOCX: Use doc_style="table" if user says "table", "tabular", "spreadsheet", "grid"
   - For PDF/DOCX: Use doc_style="paragraph" if user says "document", "essay", "article", "report", "paragraph"
   - Default for PDF/DOCX is doc_style="paragraph"

4. **Template file context**: If a template file is mentioned or previously set, infer the format from that file's extension

REASONING APPROACH:
- Think step-by-step about what the user is asking
- Consider the data type and use case
- Make intelligent inferences about the best format
- Don't default to CSV unless it truly makes sense
- Consider the user's domain (API development vs reporting vs data science)

Return JSON with: {"intent": "...", "params": {...}, "confidence": 0.0-1.0, "reasoning": "brief explanation of format choice"}"""

    context_info = context.get_context_summary()

    # Add template file context if available
    template_context = ""
    if context.user_preferences.get("template_file"):
        template_file = context.user_preferences["template_file"]
        template_format = _detect_format_from_file(template_file)
        template_context = f"\n\nTemplate file is set to: {template_file} (format: {template_format})"

    prompt = f"""Previous context:
{context_info}{template_context}

Current request: "{user_input}"

Think carefully about the user's intent and what format would best serve their needs.
Classify the intent and extract parameters using reasoning."""

    try:
        response = client.messages.create(
            model=os.getenv("SYNTH_AGENT_LLM_MODEL", "claude-sonnet-4-20250514"),
            max_tokens=1500,
            system=system_prompt,
            messages=[{"role": "user", "content": prompt}]
        )

        content = response.content[0].text

        # Extract JSON from response
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            result = json.loads(json_match.group())

            # Show reasoning if present (for transparency)
            if result.get("reasoning") and result.get("intent") == "GENERATE":
                console.print(f"[dim]💡 Reasoning: {result['reasoning']}[/dim]")

            return result
        else:
            # Fallback: Try to intelligently infer format from keywords
            return _fallback_intent_classification(user_input, context)

    except Exception as e:
        console.print(f"[yellow]⚠️  Using fallback intent classification: {e}[/yellow]")
        return _fallback_intent_classification(user_input, context)


def _fallback_intent_classification(user_input: str, context: ConversationContext) -> Dict[str, Any]:
    """Fallback intent classification using keyword matching and reasoning."""
    user_lower = user_input.lower()

    # Detect format from keywords
    format_type = "csv"  # default

    if any(word in user_lower for word in ["json", "api", "rest", "endpoint", "nested"]):
        format_type = "json"
    elif any(word in user_lower for word in ["excel", "xlsx", "spreadsheet"]):
        format_type = "excel"
    elif any(word in user_lower for word in ["pdf", "document", "paragraph", "essay", "article", "report"]) and "table" not in user_lower:
        format_type = "pdf"
    elif any(word in user_lower for word in ["word", "docx", "doc"]):
        format_type = "docx"
    elif any(word in user_lower for word in ["parquet", "big data", "analytics"]):
        format_type = "parquet"
    elif any(word in user_lower for word in ["xml", "soap", "web service"]):
        format_type = "xml"
    elif "text" in user_lower or "txt" in user_lower:
        format_type = "txt"

    # If template file is set, use its format
    if context.user_preferences.get("template_file"):
        detected = _detect_format_from_file(context.user_preferences["template_file"])
        if detected:
            format_type = detected

    # Extract row count
    rows = 10
    import re as regex
    row_match = regex.search(r'(\d+)\s*(?:rows?|records?|items?|entries)', user_lower)
    if row_match:
        rows = int(row_match.group(1))

    return {
        "intent": "GENERATE",
        "params": {
            "data_description": user_input,
            "rows": rows,
            "format": format_type
        },
        "confidence": 0.6,
        "reasoning": f"Fallback classification detected {format_type.upper()} format from keywords"
    }


async def handle_generate(client, params: Dict[str, Any], context: ConversationContext) -> str:
    """Handle data generation intent."""
    description = params.get("data_description", "data")
    rows = params.get("rows", context.user_preferences.get("default_rows", 10))
    format_type = params.get("format", context.user_preferences.get("default_format", "csv"))
    fields = params.get("fields", [])

    # Check if template file is set and should be used
    template_file = context.user_preferences.get("template_file")
    template_content = None

    if template_file and Path(template_file).exists():
        try:
            # Read template file to extract patterns
            console.print(f"[cyan]📋 Using template: {template_file}[/cyan]")
            template_content = _extract_template_content(template_file)

            # Override format based on template if not explicitly specified in params
            if "format" not in params and template_file:
                detected_format = _detect_format_from_file(template_file)
                if detected_format:
                    format_type = detected_format
                    console.print(f"[dim]Using template format: {format_type.upper()}[/dim]")
        except Exception as e:
            console.print(f"[yellow]⚠️  Could not read template file: {e}[/yellow]")
            console.print(f"[dim]Continuing with standard generation...[/dim]")

    # Unified doc_style parameter for all document formats (PDF, DOCX, Word)
    doc_style = params.get("doc_style", params.get("pdf_style", "paragraph"))
    
    # Normalize "word" to "docx"
    if format_type.lower() in ["word", "doc"]:
        format_type = "docx"
    
    # Add template context if available
    template_instruction = ""
    if template_content:
        template_instruction = f"""

TEMPLATE/PATTERN TO FOLLOW:
{template_content}

IMPORTANT: Analyze the template above and generate data that matches its structure, style, and patterns closely."""

    # Special handling for document formats (PDF, DOCX) in paragraph mode
    if format_type in ["pdf", "docx"] and doc_style == "paragraph":
        format_name = "PDF" if format_type == "pdf" else "Word"
        console.print(f"\n[cyan]📝 Generating document about {description} in {format_name} format...[/cyan]")

        # Generate actual paragraph content, not CSV rows
        generation_prompt = f"""Write a comprehensive document about: {description}

Create well-structured paragraph content with:
- Clear introduction
- Multiple detailed paragraphs with substantive information
- Professional writing style
- Proper flow and transitions
- Conclusion if appropriate
{template_instruction}

Write 3-5 paragraphs of high-quality content. Do NOT create CSV, tables, or lists. Write flowing paragraphs like an article or essay."""
    else:
        console.print(f"\n[cyan]📝 Generating {rows} rows of {description} in {format_type.upper()} format...[/cyan]")

        # Build generation prompt for structured data
        fields_text = f"\nFields to include: {', '.join(fields)}" if fields else ""
        generation_prompt = f"""Generate synthetic {description} data with {rows} rows.{fields_text}

Output format: {format_type.upper()}

Requirements:
- Make the data realistic and diverse
- Include appropriate field names and types
- Ensure data quality and consistency
{"- Output as CSV with header row" if format_type == "csv" else ""}
{"- Output as valid JSON array" if format_type == "json" else ""}
{"- Output as plain text, one item per line" if format_type == "txt" else ""}
{"- Output as valid XML document" if format_type == "xml" else ""}
{template_instruction}

Generate the data now:"""

    try:
        response = client.messages.create(
            model=os.getenv("SYNTH_AGENT_LLM_MODEL", "claude-sonnet-4-20250514"),
            max_tokens=4000,
            messages=[{"role": "user", "content": generation_prompt}]
        )
        
        generated_data = response.content[0].text
        
        # Clean up markdown formatting
        generated_data = re.sub(r'```(?:csv|json|xml|txt)?\n', '', generated_data)
        generated_data = re.sub(r'```\s*$', '', generated_data)
        generated_data = generated_data.strip()
        
        # Determine filename
        safe_desc = re.sub(r'[^\w\s-]', '', description.lower())
        safe_desc = re.sub(r'[-\s]+', '_', safe_desc)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        extensions = {
            "csv": ".csv",
            "json": ".json",
            "excel": ".xlsx",
            "parquet": ".parquet",
            "xml": ".xml",
            "txt": ".txt",
            "pdf": ".pdf",
            "docx": ".docx"
        }
        
        # Validate format
        supported_formats = list(extensions.keys())
        if format_type not in supported_formats:
            console.print(f"\n[bold red]❌ Error:[/bold red] Format '{format_type}' not supported.")
            console.print(f"[dim]Supported formats: {', '.join(supported_formats).upper()}[/dim]")
            return f"Unsupported format: {format_type}"
        
        filename = f"{safe_desc}_{timestamp}{extensions[format_type]}"
        
        # Handle Excel conversion
        if format_type == "excel":
            import pandas as pd
            from io import StringIO
            
            # First generate CSV, then convert
            temp_file = f"temp_{timestamp}.csv"
            with open(temp_file, "w", encoding="utf-8") as f:
                f.write(generated_data)
            
            df = pd.read_csv(temp_file)
            df.to_excel(filename, index=False, engine='openpyxl')
            
            # Clean up temp file
            Path(temp_file).unlink()
            
        # Handle Parquet conversion
        elif format_type == "parquet":
            import pandas as pd
            from io import StringIO
            
            temp_file = f"temp_{timestamp}.csv"
            with open(temp_file, "w", encoding="utf-8") as f:
                f.write(generated_data)
            
            df = pd.read_csv(temp_file)
            df.to_parquet(filename, index=False, engine='pyarrow')
            
            Path(temp_file).unlink()
            
        # Handle PDF conversion
        elif format_type == "pdf":
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
            import pandas as pd
            from io import StringIO
            
            # Check if user wants table format (default is paragraph)
            pdf_style = params.get("pdf_style", "paragraph")
            
            # Create PDF document
            doc = SimpleDocTemplate(filename, pagesize=letter,
                                   topMargin=inch, bottomMargin=inch,
                                   leftMargin=inch, rightMargin=inch)
            elements = []
            styles = getSampleStyleSheet()
            
            # Add title
            title = Paragraph(f"<b>{description.title()}</b>", styles['Title'])
            elements.append(title)
            elements.append(Spacer(1, 0.3 * inch))
            
            # Add metadata
            metadata_style = ParagraphStyle(
                'metadata',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey
            )
            metadata = Paragraph(
                f"<i>Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}<br/>Contains {rows} records</i>",
                metadata_style
            )
            elements.append(metadata)
            elements.append(Spacer(1, 0.4 * inch))
            
            if pdf_style == "table":
                # TABLE MODE - Parse CSV and create table
                temp_file = f"temp_{timestamp}.csv"
                with open(temp_file, "w", encoding="utf-8") as f:
                    f.write(generated_data)
                
                # Try reading CSV with different parameters to handle commas in fields
                try:
                    df = pd.read_csv(temp_file, quotechar='"', escapechar='\\')
                except:
                    try:
                        df = pd.read_csv(temp_file, sep=',', on_bad_lines='skip')
                    except:
                        # Last resort: parse manually
                        with open(temp_file, 'r', encoding='utf-8') as f:
                            lines = f.readlines()
                            if len(lines) < 2:
                                raise ValueError("Insufficient data in generated CSV")
                            # Use csv module for robust parsing
                            import csv
                            reader = csv.reader(lines)
                            data_list = list(reader)
                            df = pd.DataFrame(data_list[1:], columns=data_list[0])
                
                # Convert dataframe to table data
                table_data = [df.columns.tolist()] + df.values.tolist()
                
                # Create table with styling
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 8),
                ]))
                
                elements.append(table)
                
                # Clean up temp file
                Path(temp_file).unlink()
                
            else:
                # PARAGRAPH MODE (DEFAULT) - Natural document format
                # Create paragraph style for content
                content_style = ParagraphStyle(
                    'content',
                    parent=styles['Normal'],
                    fontSize=11,
                    leading=16,
                    alignment=TA_JUSTIFY,
                    spaceAfter=14
                )
                
                # Split generated data into paragraphs
                paragraphs = generated_data.strip().split('\n\n')
                
                # If no double newlines, try single newlines for paragraph breaks
                if len(paragraphs) == 1:
                    paragraphs = [p for p in generated_data.strip().split('\n') if p.strip()]
                
                # Add each paragraph to the PDF
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        # Clean up any special characters that might break PDF
                        para_text = para_text.replace('<', '&lt;').replace('>', '&gt;')
                        # But allow basic formatting
                        para_text = para_text.replace('&lt;b&gt;', '<b>').replace('&lt;/b&gt;', '</b>')
                        para_text = para_text.replace('&lt;i&gt;', '<i>').replace('&lt;/i&gt;', '</i>')
                        
                        paragraph = Paragraph(para_text, content_style)
                        elements.append(paragraph)
                        elements.append(Spacer(1, 0.15 * inch))
            
            # Build PDF
            doc.build(elements)
            
        # Handle DOCX (Word Document) conversion
        elif format_type == "docx":
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create Word document
            doc = Document()
            
            # Add title
            title = doc.add_heading(description.title(), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            metadata = doc.add_paragraph()
            metadata.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}").italic = True
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Blank line
            
            if doc_style == "table":
                # TABLE MODE - Parse CSV and create table
                temp_file = f"temp_{timestamp}.csv"
                with open(temp_file, "w", encoding="utf-8") as f:
                    f.write(generated_data)
                
                # Parse CSV
                import csv
                with open(temp_file, 'r', encoding='utf-8') as f:
                    csv_reader = csv.reader(f)
                    data_list = list(csv_reader)
                
                if len(data_list) > 1:
                    # Create table
                    table = doc.add_table(rows=len(data_list), cols=len(data_list[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add data to table
                    for i, row_data in enumerate(data_list):
                        row = table.rows[i]
                        for j, cell_value in enumerate(row_data):
                            row.cells[j].text = str(cell_value)
                            # Bold header row
                            if i == 0:
                                row.cells[j].paragraphs[0].runs[0].bold = True
                
                # Clean up temp file
                Path(temp_file).unlink()
            else:
                # PARAGRAPH MODE (DEFAULT) - Natural document format
                paragraphs = generated_data.strip().split('\n\n')
                
                # If no double newlines, try single newlines
                if len(paragraphs) == 1:
                    paragraphs = [p for p in generated_data.strip().split('\n') if p.strip()]
                
                # Add each paragraph to the Word document
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        p = doc.add_paragraph(para_text)
                        # Format paragraph
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in p.runs:
                            run.font.size = Pt(11)
            
            # Save Word document
            doc.save(filename)
            
        else:
            # Save directly (CSV, JSON, XML, TXT)
            with open(filename, "w", encoding="utf-8") as f:
                f.write(generated_data)
        
        # Update context
        context.last_generated_file = filename
        
        file_size = Path(filename).stat().st_size
        console.print(f"\n[bold green]✅ Success![/bold green] Generated [cyan]{filename}[/cyan] ({file_size:,} bytes)")
        
        # Different message for document formats vs data formats
        if format_type in ["pdf", "docx"] and doc_style == "paragraph":
            format_name = "PDF" if format_type == "pdf" else "Word"
            console.print(f"[dim]Document generated in {format_name} paragraph format[/dim]")
        else:
            console.print(f"[dim]Contains {rows} rows in {format_type.upper()} format[/dim]")
        
        return f"Successfully generated {filename}"
        
    except Exception as e:
        error_msg = f"Failed to generate data: {str(e)}"
        console.print(f"\n[bold red]❌ Error:[/bold red] {error_msg}")
        return error_msg


async def handle_analyze(client, params: Dict[str, Any], context: ConversationContext) -> str:
    """Handle data analysis intent."""
    filename = params.get("filename", context.last_generated_file)
    
    if not filename:
        return "❌ Please specify a file to analyze, or generate data first."
    
    if not Path(filename).exists():
        return f"❌ File '{filename}' not found."
    
    console.print(f"\n[cyan]🔍 Analyzing {filename}...[/cyan]")
    
    try:
        # Read file content
        content = Path(filename).read_text(encoding="utf-8")
        
        # Limit content for analysis
        if len(content) > 5000:
            content = content[:5000] + "\n... (truncated)"
        
        analysis_prompt = f"""Analyze this data file and provide insights:

Filename: {filename}
Content:
{content}

Provide:
1. Data structure overview (columns, types, format)
2. Key statistics (row count, unique values, ranges)
3. Data quality assessment (missing values, anomalies)
4. Patterns and insights
5. Recommendations for use

Be concise and actionable."""

        response = client.messages.create(
            model=os.getenv("SYNTH_AGENT_LLM_MODEL", "claude-sonnet-4-20250514"),
            max_tokens=2000,
            messages=[{"role": "user", "content": analysis_prompt}]
        )
        
        analysis = response.content[0].text
        
        console.print(Panel(Markdown(analysis), title=f"📊 Analysis: {filename}", border_style="cyan"))
        
        return f"Analyzed {filename}"
        
    except Exception as e:
        error_msg = f"Failed to analyze file: {str(e)}"
        console.print(f"\n[bold red]❌ Error:[/bold red] {error_msg}")
        return error_msg


def handle_configure(params: Dict[str, Any], context: ConversationContext) -> str:
    """Handle configuration changes."""
    setting = params.get("setting_name", "")
    value = params.get("value")

    if "format" in setting.lower():
        if value in ["csv", "json", "excel", "parquet", "xml", "txt"]:
            context.user_preferences["default_format"] = value
            console.print(f"\n[green]✅ Default format set to {value.upper()}[/green]")
            return f"Updated default format to {value}"
        else:
            console.print(f"\n[red]❌ Invalid format: {value}[/red]")
            return "Invalid format"

    elif "rows" in setting.lower():
        try:
            rows = int(value)
            context.user_preferences["default_rows"] = rows
            console.print(f"\n[green]✅ Default rows set to {rows}[/green]")
            return f"Updated default rows to {rows}"
        except:
            console.print(f"\n[red]❌ Invalid number: {value}[/red]")
            return "Invalid number"

    elif "template" in setting.lower() or "pattern" in setting.lower():
        # Handle template file setting
        file_path = str(value).strip().strip('"').strip("'")

        # Check if file exists
        if Path(file_path).exists():
            context.user_preferences["template_file"] = file_path

            # Detect format from file extension
            detected_format = _detect_format_from_file(file_path)
            if detected_format:
                context.user_preferences["default_format"] = detected_format
                console.print(f"\n[green]✅ Template file set to: {file_path}[/green]")
                console.print(f"[cyan]📝 Auto-detected format: {detected_format.upper()}[/cyan]")
                return f"Updated template file to {file_path} (format: {detected_format})"
            else:
                console.print(f"\n[green]✅ Template file set to: {file_path}[/green]")
                console.print(f"[yellow]⚠️  Could not auto-detect format. Please set format manually if needed.[/yellow]")
                return f"Updated template file to {file_path}"
        else:
            console.print(f"\n[red]❌ File not found: {file_path}[/red]")
            console.print(f"[dim]Please check the path and try again.[/dim]")
            return f"File not found: {file_path}"

    else:
        console.print(f"\n[yellow]⚠️  Unknown setting: {setting}[/yellow]")
        console.print(f"[dim]Available settings: format, rows, template_file[/dim]")
        return f"Unknown setting: {setting}"


def _detect_format_from_file(file_path: str) -> Optional[str]:
    """Detect output format from template file extension."""
    ext = Path(file_path).suffix.lower()

    format_map = {
        ".csv": "csv",
        ".json": "json",
        ".xlsx": "excel",
        ".xls": "excel",
        ".parquet": "parquet",
        ".xml": "xml",
        ".txt": "txt",
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx"
    }

    return format_map.get(ext)


def _extract_template_content(file_path: str) -> str:
    """Extract content from template file to understand patterns."""
    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".csv":
            with open(file_path, 'r', encoding='utf-8') as f:
                # Read first few lines to get structure
                lines = [f.readline() for _ in range(min(5, sum(1 for _ in f) + 1))]
                return "".join(lines)

        elif ext == ".json":
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                # Limit content size
                return content[:2000] if len(content) > 2000 else content

        elif ext in [".xlsx", ".xls"]:
            import pandas as pd
            df = pd.read_excel(file_path, nrows=5)
            return df.to_string()

        elif ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                # Read first 2000 characters
                return f.read(2000)

        elif ext == ".xml":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(2000)

        elif ext == ".pdf":
            # For PDF, extract text from first page
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                if len(reader.pages) > 0:
                    text = reader.pages[0].extract_text()
                    return text[:2000] if len(text) > 2000 else text
            except:
                return "PDF content (could not extract text)"

        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(file_path)
                # Extract first few paragraphs
                paragraphs = [p.text for p in doc.paragraphs[:5]]
                return "\n".join(paragraphs)
            except:
                return "Word document content (could not extract text)"

        else:
            # Generic text extraction
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(2000)

    except Exception as e:
        return f"Could not extract template content: {str(e)}"


def handle_file_ops(params: Dict[str, Any], context: ConversationContext) -> str:
    """Handle file operations."""
    operation = params.get("operation", "list")
    filename = params.get("filename")
    
    if operation == "list":
        # List all generated files
        patterns = ["*.csv", "*.json", "*.xlsx", "*.parquet", "*.xml", "*.txt"]
        files = []
        for pattern in patterns:
            files.extend(Path(".").glob(pattern))
        
        if not files:
            console.print("\n[yellow]No data files found in current directory.[/yellow]")
            return "No files found"
        
        table = Table(title="📁 Your Data Files", show_header=True, header_style="bold cyan")
        table.add_column("Filename", style="green")
        table.add_column("Size", style="yellow")
        table.add_column("Modified", style="dim")
        
        for file in sorted(files, key=lambda f: f.stat().st_mtime, reverse=True):
            size = file.stat().st_size
            modified = datetime.fromtimestamp(file.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
            
            size_str = f"{size:,} bytes" if size < 1024 else f"{size/1024:.1f} KB"
            table.add_row(file.name, size_str, modified)
        
        console.print("\n")
        console.print(table)
        console.print()
        
        return f"Listed {len(files)} files"
    
    elif operation == "delete":
        if not filename:
            console.print("\n[red]❌ Please specify a filename to delete.[/red]")
            return "No filename specified"
        
        file_path = Path(filename)
        if file_path.exists():
            file_path.unlink()
            console.print(f"\n[green]✅ Deleted {filename}[/green]")
            return f"Deleted {filename}"
        else:
            console.print(f"\n[red]❌ File '{filename}' not found.[/red]")
            return f"File not found: {filename}"
    
    else:
        console.print(f"\n[yellow]⚠️  Unknown operation: {operation}[/yellow]")
        return f"Unknown operation: {operation}"


def handle_help() -> str:
    """Show help information."""
    print_examples()
    
    help_panel = """
## 💡 Quick Tips

**Generating Data:**
- Be specific about what you need: "customer data with emails and phone numbers"
- Specify quantity: "100 rows", "50 records", "20 items"
- Choose format: "as JSON", "in Excel format", "CSV file"

**File Formats:**
- **CSV** - Best for spreadsheets, simple tables
- **JSON** - Best for APIs, nested data structures
- **Excel** - Best for business reports, formatted output
- **Parquet** - Best for big data, analytics pipelines
- **XML** - Best for legacy systems, web services
- **TXT** - Best for lists, documents, questions

**Analysis:**
- "analyze my data" - Analyze last generated file
- "analyze customers.csv" - Analyze specific file
- "show patterns in sales.json" - Deep pattern analysis

**Settings:**
- "use JSON by default" - Change default format
- "make default 100 rows" - Change default row count
- "set template to customers.csv" - Use a file as pattern/template
- "use data.json as template" - Pattern-based generation

**Template Files:**
The system can use existing files as templates to match structure and patterns:
- Automatically detects format from template file extension
- Extracts patterns from CSV, JSON, Excel, PDF, Word, XML, and more
- Generates new data matching the template's structure and style

**File Management:**
- "show my files" - List all data files
- "delete old_data.csv" - Remove specific file
"""
    
    console.print(Panel(Markdown(help_panel), title="📚 Help & Tips", border_style="yellow"))
    return "Showed help"


async def handle_question(client, user_input: str) -> str:
    """Handle general questions using Claude."""
    console.print("\n[cyan]🤔 Let me think about that...[/cyan]")
    
    system_prompt = """You are a helpful assistant for a synthetic data generation tool.
Answer questions about:
- Data generation techniques
- File formats and their uses
- Best practices for synthetic data
- How to use the tool
- Data quality and patterns

Be concise, helpful, and practical."""

    try:
        response = client.messages.create(
            model=os.getenv("SYNTH_AGENT_LLM_MODEL", "claude-sonnet-4-20250514"),
            max_tokens=1500,
            system=system_prompt,
            messages=[{"role": "user", "content": user_input}]
        )
        
        answer = response.content[0].text
        console.print(Panel(Markdown(answer), title="💡 Answer", border_style="blue"))
        
        return "Answered question"
    except Exception as e:
        console.print(f"\n[red]❌ Error: {str(e)}[/red]")
        return f"Error: {str(e)}"


@app.command()
def start():
    """
    🤖 Start the fully NLP-based chat interface.
    
    Interact with the tool using natural language - no commands to remember!
    """
    asyncio.run(main_chat())


async def main_chat():
    """Main chat loop."""
    try:
        import anthropic
        
        # Load environment
        load_dotenv()
        
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            console.print("[bold red]❌ ANTHROPIC_API_KEY not found in environment variables[/bold red]")
            console.print("\n[yellow]Please set your API key:[/yellow]")
            console.print("1. Create a .env file in the project directory")
            console.print("2. Add: ANTHROPIC_API_KEY=your-key-here")
            raise typer.Exit(1)
        
        # Create client and context
        client = anthropic.Anthropic(api_key=api_key)
        context = ConversationContext()
        
        # Show welcome
        print_welcome()
        
        console.print("\n[dim]💬 Start chatting! Type 'help' for examples, 'exit' to quit.[/dim]\n")
        
        # Main conversation loop
        while True:
            try:
                # Get user input
                user_input = Prompt.ask("\n[bold green]You[/bold green]").strip()
                
                if not user_input:
                    continue
                
                # Check for exit
                if user_input.lower() in ["exit", "quit", "bye", "goodbye", "done"]:
                    console.print("\n[cyan]👋 Goodbye! Happy data generating![/cyan]\n")
                    break
                
                # Add to context
                context.add_message("user", user_input)
                
                # Special cases
                if user_input.lower() in ["help", "examples", "show examples"]:
                    handle_help()
                    continue
                
                # Classify intent
                console.print("[dim]🤔 Understanding...[/dim]", end="")
                intent_result = await classify_intent(client, user_input, context)
                console.print("\r" + " " * 50 + "\r", end="")  # Clear line
                
                intent = intent_result.get("intent", "CHAT")
                params = intent_result.get("params", {})
                
                # Handle based on intent
                result = ""
                if intent == "GENERATE":
                    result = await handle_generate(client, params, context)
                elif intent == "ANALYZE":
                    result = await handle_analyze(client, params, context)
                elif intent == "CONFIGURE":
                    result = handle_configure(params, context)
                elif intent == "FILE_OPS":
                    result = handle_file_ops(params, context)
                elif intent == "HELP":
                    result = handle_help()
                elif intent in ["QUESTION", "CHAT"]:
                    result = await handle_question(client, user_input)
                else:
                    # Fallback to question handling
                    result = await handle_question(client, user_input)
                
                # Add result to context
                context.add_message("assistant", result)
                
            except KeyboardInterrupt:
                console.print("\n\n[cyan]👋 Interrupted. Goodbye![/cyan]\n")
                break
            except Exception as e:
                console.print(f"\n[red]❌ Error: {str(e)}[/red]")
                console.print("[dim]Type 'help' for assistance or 'exit' to quit.[/dim]")
    
    except Exception as e:
        console.print(f"\n[bold red]❌ Fatal Error:[/bold red] {str(e)}")
        raise typer.Exit(1)


@app.callback(invoke_without_command=True)
def main(ctx: typer.Context):
    """
    🤖 Synthetic Data Generator - Fully NLP Chat Interface
    
    Just run 'synth-agent' or 'synth-agent start' to begin chatting!
    """
    if ctx.invoked_subcommand is None:
        # No subcommand - start chat directly
        asyncio.run(main_chat())


if __name__ == "__main__":
    app()
